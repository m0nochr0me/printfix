"""Table fixes: auto-fit and text resize."""

from __future__ import annotations

import asyncio

from lxml import etree  # type: ignore

from app.core.log import logger
from app.schema.fix import FixResult

__all__ = ("auto_fit_tables", "resize_table_text")

_WP_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


async def auto_fit_tables(file_path: str, job_id: str) -> FixResult:
    """Set all tables to auto-fit within page width in a DOCX."""
    return await asyncio.to_thread(_auto_fit_tables_sync, file_path, job_id)


def _auto_fit_tables_sync(file_path: str, job_id: str) -> FixResult:
    from docx import Document
    from docx.shared import Emu

    doc = Document(file_path)
    fixed = 0

    # Get printable width from first section
    if not doc.sections:
        return FixResult(
            tool_name="auto_fit_tables",
            job_id=job_id,
            success=False,
            description="No sections found in document",
            error="No sections",
        )

    section = doc.sections[0]
    printable_width = (
        (section.page_width or 0)
        - (section.left_margin or 0)
        - (section.right_margin or 0)
    )

    for table in doc.tables:
        tbl = table._tbl

        # Set table width to 100% of page width
        tbl_pr = tbl.find(f"{{{_WP_NS}}}tblPr")
        if tbl_pr is None:
            tbl_pr = etree.SubElement(tbl, f"{{{_WP_NS}}}tblPr")

        # Set preferred width to page width
        tbl_w = tbl_pr.find(f"{{{_WP_NS}}}tblW")
        if tbl_w is None:
            tbl_w = etree.SubElement(tbl_pr, f"{{{_WP_NS}}}tblW")
        tbl_w.set(f"{{{_WP_NS}}}w", str(int(printable_width)))
        tbl_w.set(f"{{{_WP_NS}}}type", "dxa")

        # Set autofit layout
        tbl_layout = tbl_pr.find(f"{{{_WP_NS}}}tblLayout")
        if tbl_layout is None:
            tbl_layout = etree.SubElement(tbl_pr, f"{{{_WP_NS}}}tblLayout")
        tbl_layout.set(f"{{{_WP_NS}}}type", "autofit")

        # Scale column widths proportionally to fit
        _scale_columns(table, printable_width)
        fixed += 1

    doc.save(file_path)
    return FixResult(
        tool_name="auto_fit_tables",
        job_id=job_id,
        success=True,
        description=f"Set {fixed} table(s) to auto-fit within page width",
        after_value=f"{fixed} tables auto-fitted",
    )


def _scale_columns(table, target_width: int) -> None:
    """Scale column widths proportionally to fit target width."""
    from docx.shared import Emu

    if not table.rows:
        return

    # Measure current total width from first row
    current_widths: list[int] = []
    for cell in table.rows[0].cells:
        current_widths.append(cell.width or 0)

    total = sum(current_widths)
    if total <= 0 or total <= target_width:
        return

    # Scale each column proportionally
    scale = target_width / total
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(current_widths) and current_widths[i] > 0:
                cell.width = int(current_widths[i] * scale)


async def resize_table_text(
    file_path: str,
    job_id: str,
    table_index: int = 0,
    max_font_size_pt: float = 10.0,
) -> FixResult:
    """Reduce font size in a specific table's cells (0-indexed)."""
    return await asyncio.to_thread(
        _resize_table_text_sync,
        file_path,
        job_id,
        table_index,
        max_font_size_pt,
    )


def _resize_table_text_sync(
    file_path: str,
    job_id: str,
    table_index: int,
    max_font_size_pt: float,
) -> FixResult:
    from docx import Document
    from docx.shared import Pt

    doc = Document(file_path)

    if table_index >= len(doc.tables):
        return FixResult(
            tool_name="resize_table_text",
            job_id=job_id,
            success=False,
            description=f"Table index {table_index} out of range (document has {len(doc.tables)} tables)",
            error="Table index out of range",
        )

    table = doc.tables[table_index]
    adjusted = 0

    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.font.size is not None:
                        pt = run.font.size / 12700
                        if pt > max_font_size_pt:
                            run.font.size = Pt(max_font_size_pt)
                            adjusted += 1

    doc.save(file_path)
    return FixResult(
        tool_name="resize_table_text",
        job_id=job_id,
        success=True,
        description=f"Capped font size to {max_font_size_pt}pt in table {table_index} ({adjusted} run(s) adjusted)",
        after_value=f"max {max_font_size_pt}pt",
    )
