"""Page layout fixes: margins, page size, orientation, blank page removal."""

from __future__ import annotations

import asyncio

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches

from app.schema.fix import FixResult

__all__ = (
    "adjust_paragraph_indents",
    "remove_blank_pages",
    "set_margins",
    "set_orientation",
    "set_page_size",
)

_WP_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_EMU_PER_INCH = 914400


async def set_margins(
    file_path: str,
    job_id: str,
    top: float = 0.5,
    bottom: float = 0.5,
    left: float = 1.0,
    right: float = 0.75,
) -> FixResult:
    """Set margins on all sections of a DOCX. Values in inches."""
    return await asyncio.to_thread(
        _set_margins_sync,
        file_path,
        job_id,
        top,
        bottom,
        left,
        right,
    )


def _set_margins_sync(
    file_path: str,
    job_id: str,
    top: float,
    bottom: float,
    left: float,
    right: float,
) -> FixResult:

    doc = Document(file_path)
    changed_sections = 0
    old_values: list[str] = []

    for i, section in enumerate(doc.sections, 1):
        old = (
            f"S{i}: T={_emu_to_in(section.top_margin)} "
            f"B={_emu_to_in(section.bottom_margin)} "
            f"L={_emu_to_in(section.left_margin)} "
            f"R={_emu_to_in(section.right_margin)}"
        )
        old_values.append(old)

        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)
        changed_sections += 1

    doc.save(file_path)
    return FixResult(
        tool_name="set_margins",
        job_id=job_id,
        success=True,
        description=f'Set margins to T={top}" B={bottom}" L={left}" R={right}" on {changed_sections} section(s)',
        before_value="; ".join(old_values),
        after_value=f'T={top}" B={bottom}" L={left}" R={right}"',
    )


async def set_page_size(
    file_path: str,
    job_id: str,
    width: float = 8.5,
    height: float = 11.0,
) -> FixResult:
    """Set page size on all sections. Values in inches. A4=8.27x11.69, Letter=8.5x11."""
    return await asyncio.to_thread(_set_page_size_sync, file_path, job_id, width, height)


def _set_page_size_sync(
    file_path: str,
    job_id: str,
    width: float,
    height: float,
) -> FixResult:

    doc = Document(file_path)
    changed = 0

    for section in doc.sections:
        section.page_width = Inches(width)
        section.page_height = Inches(height)
        changed += 1

    doc.save(file_path)
    return FixResult(
        tool_name="set_page_size",
        job_id=job_id,
        success=True,
        description=f'Set page size to {width}"x{height}" on {changed} section(s)',
        after_value=f'{width}"x{height}"',
    )


async def set_orientation(
    file_path: str,
    job_id: str,
    orientation: str = "portrait",
) -> FixResult:
    """Set orientation on all sections. Swaps width/height as needed."""
    return await asyncio.to_thread(_set_orientation_sync, file_path, job_id, orientation)


def _set_orientation_sync(
    file_path: str,
    job_id: str,
    orientation: str,
) -> FixResult:

    doc = Document(file_path)
    target = WD_ORIENT.PORTRAIT if orientation == "portrait" else WD_ORIENT.LANDSCAPE
    changed = 0

    for section in doc.sections:
        current = section.orientation
        if current != target:
            # Swap dimensions
            w, h = section.page_width, section.page_height
            section.page_width = h
            section.page_height = w
            section.orientation = target
            changed += 1

    doc.save(file_path)
    return FixResult(
        tool_name="set_orientation",
        job_id=job_id,
        success=True,
        description=f"Set orientation to {orientation} ({changed} section(s) changed)",
        after_value=orientation,
    )


async def remove_blank_pages(file_path: str, job_id: str) -> FixResult:
    """Remove consecutive page breaks that create blank pages in DOCX."""
    return await asyncio.to_thread(_remove_blank_pages_sync, file_path, job_id)


def _remove_blank_pages_sync(file_path: str, job_id: str) -> FixResult:

    doc = Document(file_path)
    removed = 0
    prev_was_break = False

    for para in list(doc.paragraphs):
        is_break_only = _is_page_break_paragraph(para)

        if is_break_only and prev_was_break:
            # Remove this paragraph (it creates a blank page)
            parent = para._element.getparent()
            if parent is not None:
                parent.remove(para._element)
                removed += 1
                continue

        prev_was_break = is_break_only

    doc.save(file_path)
    return FixResult(
        tool_name="remove_blank_pages",
        job_id=job_id,
        success=True,
        description=f"Removed {removed} consecutive page break(s) creating blank pages",
        after_value=f"{removed} removed",
    )


async def adjust_paragraph_indents(
    file_path: str,
    job_id: str,
    left_inches: float = 0.5,
    right_inches: float = 0.5,
    first_line_inches: float = 0.5,
    strategy: str = "cap",
) -> FixResult:
    """Adjust paragraph indents in a DOCX. Strategy: 'cap' or 'scale'.

    - cap: clamp each indent at the specified maximum
    - scale: proportionally shrink all indents so the largest equals the max
    """
    return await asyncio.to_thread(
        _adjust_paragraph_indents_sync,
        file_path,
        job_id,
        left_inches,
        right_inches,
        first_line_inches,
        strategy,
    )


def _adjust_paragraph_indents_sync(
    file_path: str,
    job_id: str,
    left_inches: float,
    right_inches: float,
    first_line_inches: float,
    strategy: str,
) -> FixResult:

    doc = Document(file_path)
    max_left = Inches(left_inches)
    max_right = Inches(right_inches)
    max_first = Inches(first_line_inches)

    # Collect all paragraphs (body + table cells)
    all_paras: list = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paras.extend(cell.paragraphs)

    scale_first, scale_left, scale_right = 1.0, 1.0, 1.0

    if strategy == "scale":
        # Find the largest indent values to compute scale factors
        largest_left = 0
        largest_right = 0
        largest_first = 0
        for para in all_paras:
            pf = para.paragraph_format
            largest_left = max(largest_left, pf.left_indent or 0)
            largest_right = max(largest_right, pf.right_indent or 0)
            fl = pf.first_line_indent or 0
            if fl > 0:
                largest_first = max(largest_first, fl)

        scale_left = max_left / largest_left if largest_left > max_left else 1.0
        scale_right = max_right / largest_right if largest_right > max_right else 1.0
        scale_first = max_first / largest_first if largest_first > max_first else 1.0

    # Get margin info for negative indent clamping
    min_neg_left = 0
    min_neg_right = 0
    if doc.sections:
        sec = doc.sections[0]
        # Allow negative indent up to half the margin (leave some safe space)
        lm = sec.left_margin or 0
        rm = sec.right_margin or 0
        min_neg_left = -(lm // 2)  # e.g. margin 0.5" → allow -0.25" indent
        min_neg_right = -(rm // 2)

    adjusted = 0
    old_max_left = 0
    old_min_left = 0
    old_max_right = 0
    old_min_right = 0

    for para in all_paras:
        pf = para.paragraph_format
        left = pf.left_indent or 0
        right = pf.right_indent or 0
        first_line = pf.first_line_indent or 0
        changed = False

        old_max_left = max(old_max_left, left)
        old_min_left = min(old_min_left, left)
        old_max_right = max(old_max_right, right)
        old_min_right = min(old_min_right, right)

        if strategy == "scale":
            if left > max_left:
                pf.left_indent = int(left * scale_left)
                changed = True
            if right > max_right:
                pf.right_indent = int(right * scale_right)
                changed = True
            if first_line > 0 and first_line > max_first:
                pf.first_line_indent = int(first_line * scale_first)
                changed = True
            # Clamp negative indents even in scale mode
            if left < min_neg_left:
                pf.left_indent = min_neg_left
                changed = True
            if right < min_neg_right:
                pf.right_indent = min_neg_right
                changed = True
        else:  # cap
            # Cap positive indents
            if left > max_left:
                pf.left_indent = max_left
                changed = True
            if right > max_right:
                pf.right_indent = max_right
                changed = True
            if first_line > max_first:
                pf.first_line_indent = max_first
                changed = True
            # Cap negative indents that push content into/past margins
            if left < min_neg_left:
                pf.left_indent = min_neg_left
                changed = True
            if right < min_neg_right:
                pf.right_indent = min_neg_right
                changed = True
            # Cap negative first-line (hanging) indents
            if first_line < -max_first:
                pf.first_line_indent = -max_first
                changed = True

        if changed:
            adjusted += 1

    doc.save(file_path)

    old_max_l = old_max_left / _EMU_PER_INCH
    old_min_l = old_min_left / _EMU_PER_INCH
    old_max_r = old_max_right / _EMU_PER_INCH
    old_min_r = old_min_right / _EMU_PER_INCH

    before_parts = [f'L={old_max_l:.2f}"']
    if old_min_left < 0:
        before_parts.append(f'L(neg)={old_min_l:+.2f}"')
    before_parts.append(f'R={old_max_r:.2f}"')
    if old_min_right < 0:
        before_parts.append(f'R(neg)={old_min_r:+.2f}"')

    return FixResult(
        tool_name="adjust_paragraph_indents",
        job_id=job_id,
        success=True,
        description=(
            f"Adjusted indents on {adjusted} paragraph(s) using '{strategy}' strategy"
            f' (max L={left_inches}" R={right_inches}" FL={first_line_inches}")'
        ),
        before_value="max " + " ".join(before_parts),
        after_value=f'cap L={left_inches}" R={right_inches}" FL={first_line_inches}"',
    )


def _is_page_break_paragraph(para) -> bool:
    """Check if a paragraph is only a page break with no real content."""
    text = para.text.strip()
    if text:
        return False

    has_break = False
    for run in para.runs:
        xml = run._element.xml
        if 'w:type="page"' in xml or "w:type='page'" in xml:
            has_break = True

    if not has_break and para.paragraph_format.page_break_before:
        has_break = True

    return has_break


def _emu_to_in(emu) -> str:
    """Convert EMU to inches string for display."""
    if emu is None:
        return "?"
    return f'{emu / 914400:.2f}"'
