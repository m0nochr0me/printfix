"""Font fixes: replacement, size adjustment, spacing, and widow/orphan control."""

from __future__ import annotations

import asyncio

from app.schema.fix import FixResult

__all__ = (
    "adjust_font_size",
    "normalize_paragraph_spacing",
    "replace_font",
    "set_line_spacing",
    "set_widow_orphan_control",
)


async def replace_font(
    file_path: str,
    job_id: str,
    from_font: str,
    to_font: str,
) -> FixResult:
    """Replace all occurrences of a font with another in a DOCX."""
    return await asyncio.to_thread(
        _replace_font_sync,
        file_path,
        job_id,
        from_font,
        to_font,
    )


def _replace_font_sync(
    file_path: str,
    job_id: str,
    from_font: str,
    to_font: str,
) -> FixResult:
    from docx import Document

    doc = Document(file_path)
    replaced = 0

    # Walk paragraphs
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.name == from_font:
                run.font.name = to_font
                replaced += 1

    # Walk table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.font.name == from_font:
                            run.font.name = to_font
                            replaced += 1

    doc.save(file_path)
    return FixResult(
        tool_name="replace_font",
        job_id=job_id,
        success=True,
        description=f"Replaced '{from_font}' with '{to_font}' in {replaced} run(s)",
        before_value=from_font,
        after_value=to_font,
    )


async def adjust_font_size(
    file_path: str,
    job_id: str,
    min_size_pt: float | None = None,
    max_size_pt: float | None = None,
) -> FixResult:
    """Clamp all font sizes to a min/max range in a DOCX."""
    return await asyncio.to_thread(
        _adjust_font_size_sync,
        file_path,
        job_id,
        min_size_pt,
        max_size_pt,
    )


def _adjust_font_size_sync(
    file_path: str,
    job_id: str,
    min_size_pt: float | None,
    max_size_pt: float | None,
) -> FixResult:
    from docx import Document
    from docx.shared import Pt

    doc = Document(file_path)
    adjusted = 0

    def clamp_run(run) -> bool:
        nonlocal adjusted
        if run.font.size is None:
            return False

        pt = run.font.size / 12700  # EMU to pt
        new_pt = pt

        if min_size_pt is not None and pt < min_size_pt:
            new_pt = min_size_pt
        if max_size_pt is not None and pt > max_size_pt:
            new_pt = max_size_pt

        if new_pt != pt:
            run.font.size = Pt(new_pt)
            adjusted += 1
            return True
        return False

    for para in doc.paragraphs:
        for run in para.runs:
            clamp_run(run)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        clamp_run(run)

    doc.save(file_path)

    range_desc = ""
    if min_size_pt is not None and max_size_pt is not None:
        range_desc = f"{min_size_pt}-{max_size_pt}pt"
    elif min_size_pt is not None:
        range_desc = f">={min_size_pt}pt"
    elif max_size_pt is not None:
        range_desc = f"<={max_size_pt}pt"

    return FixResult(
        tool_name="adjust_font_size",
        job_id=job_id,
        success=True,
        description=f"Adjusted {adjusted} run(s) to {range_desc}",
        after_value=range_desc,
    )


# ── Widow / Orphan control ──────────────────────────────────────────────

_WP_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


async def set_widow_orphan_control(
    file_path: str,
    job_id: str,
    enable: bool = True,
) -> FixResult:
    """Enable or disable widow/orphan control on all paragraphs in a DOCX.

    When enabled, Word prevents single lines from being stranded at the
    top (widow) or bottom (orphan) of a page.
    """
    return await asyncio.to_thread(
        _set_widow_orphan_control_sync,
        file_path,
        job_id,
        enable,
    )


def _set_widow_orphan_control_sync(
    file_path: str,
    job_id: str,
    enable: bool,
) -> FixResult:
    from lxml import etree

    from docx import Document

    doc = Document(file_path)
    changed = 0

    val = "1" if enable else "0"

    for para in doc.paragraphs:
        ppr = para._element.find(f"{{{_WP_NS}}}pPr")
        if ppr is None:
            ppr = etree.SubElement(para._element, f"{{{_WP_NS}}}pPr")
            para._element.insert(0, ppr)

        wc = ppr.find(f"{{{_WP_NS}}}widowControl")
        current_val = wc.get(f"{{{_WP_NS}}}val") if wc is not None else None

        if wc is None:
            if enable:  # Default in Word is enabled, only add if we need to set
                wc = etree.SubElement(ppr, f"{{{_WP_NS}}}widowControl")
                wc.set(f"{{{_WP_NS}}}val", val)
                changed += 1
        elif current_val != val:
            wc.set(f"{{{_WP_NS}}}val", val)
            changed += 1

    doc.save(file_path)

    action = "Enabled" if enable else "Disabled"
    return FixResult(
        tool_name="set_widow_orphan_control",
        job_id=job_id,
        success=True,
        description=f"{action} widow/orphan control on {changed} paragraph(s)",
        after_value=f"widowControl={'on' if enable else 'off'}",
    )


# ── Paragraph spacing normalization ─────────────────────────────────────


async def normalize_paragraph_spacing(
    file_path: str,
    job_id: str,
    before_pt: float = 0.0,
    after_pt: float = 8.0,
) -> FixResult:
    """Standardize space-before and space-after on body paragraphs.

    Skips heading styles to preserve document hierarchy.
    Values in points.
    """
    return await asyncio.to_thread(
        _normalize_paragraph_spacing_sync,
        file_path,
        job_id,
        before_pt,
        after_pt,
    )


def _normalize_paragraph_spacing_sync(
    file_path: str,
    job_id: str,
    before_pt: float,
    after_pt: float,
) -> FixResult:
    from docx import Document
    from docx.shared import Pt

    doc = Document(file_path)
    changed = 0

    for para in doc.paragraphs:
        # Skip headings — preserve hierarchy
        style_name = (para.style.name or "").lower() if para.style else ""
        if style_name.startswith("heading") or style_name.startswith("title"):
            continue

        pf = para.paragraph_format
        did_change = False

        if pf.space_before is not None and pf.space_before != Pt(before_pt):
            pf.space_before = Pt(before_pt)
            did_change = True
        elif pf.space_before is None and before_pt > 0:
            pf.space_before = Pt(before_pt)
            did_change = True

        if pf.space_after is not None and pf.space_after != Pt(after_pt):
            pf.space_after = Pt(after_pt)
            did_change = True
        elif pf.space_after is None and after_pt != 8.0:
            # Default is typically 8pt, only set if non-default
            pf.space_after = Pt(after_pt)
            did_change = True

        if did_change:
            changed += 1

    doc.save(file_path)

    return FixResult(
        tool_name="normalize_paragraph_spacing",
        job_id=job_id,
        success=True,
        description=(f"Normalized spacing on {changed} paragraph(s): before={before_pt}pt, after={after_pt}pt"),
        after_value=f"before={before_pt}pt after={after_pt}pt",
    )


# ── Line spacing ────────────────────────────────────────────────────────


async def set_line_spacing(
    file_path: str,
    job_id: str,
    spacing: float = 1.15,
    rule: str = "multiple",
) -> FixResult:
    """Force consistent line spacing across all body paragraphs.

    Args:
        spacing: The line spacing value. For 'multiple' rule: 1.0=single,
            1.15, 1.5, 2.0=double. For 'exact'/'at_least': value in points.
        rule: 'multiple', 'exact', or 'at_least'
    """
    return await asyncio.to_thread(
        _set_line_spacing_sync,
        file_path,
        job_id,
        spacing,
        rule,
    )


def _set_line_spacing_sync(
    file_path: str,
    job_id: str,
    spacing: float,
    rule: str,
) -> FixResult:
    from docx import Document
    from docx.enum.text import WD_LINE_SPACING
    from docx.shared import Pt

    rule_map = {
        "multiple": WD_LINE_SPACING.MULTIPLE,
        "exact": WD_LINE_SPACING.EXACTLY,
        "at_least": WD_LINE_SPACING.AT_LEAST,
        "single": WD_LINE_SPACING.SINGLE,
        "double": WD_LINE_SPACING.DOUBLE,
        "one_point_five": WD_LINE_SPACING.ONE_POINT_FIVE,
    }

    wd_rule = rule_map.get(rule, WD_LINE_SPACING.MULTIPLE)
    doc = Document(file_path)
    changed = 0

    for para in doc.paragraphs:
        # Skip headings
        style_name = (para.style.name or "").lower() if para.style else ""
        if style_name.startswith("heading") or style_name.startswith("title"):
            continue

        pf = para.paragraph_format

        if rule == "multiple":
            if pf.line_spacing != spacing or pf.line_spacing_rule != wd_rule:
                pf.line_spacing = spacing
                pf.line_spacing_rule = wd_rule
                changed += 1
        elif rule in ("exact", "at_least"):
            target = Pt(spacing)
            if pf.line_spacing != target or pf.line_spacing_rule != wd_rule:
                pf.line_spacing = target
                pf.line_spacing_rule = wd_rule
                changed += 1
        else:
            # single, double, one_point_five — rule-only, no numeric value
            if pf.line_spacing_rule != wd_rule:
                pf.line_spacing_rule = wd_rule
                changed += 1

    doc.save(file_path)

    return FixResult(
        tool_name="set_line_spacing",
        job_id=job_id,
        success=True,
        description=f"Set line spacing to {spacing} ({rule}) on {changed} paragraph(s)",
        after_value=f"{spacing} ({rule})",
    )
