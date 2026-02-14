"""DOCX document cleanup fixes: tracked changes, hidden text, empty paragraphs."""

import asyncio

from docx import Document
from docx.shared import Pt

from app.schema.fix import FixResult

__all__ = (
    "accept_tracked_changes",
    "normalize_styles",
    "remove_empty_paragraphs",
    "strip_hidden_text",
)

_WP_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


async def accept_tracked_changes(
    file_path: str,
    job_id: str,
) -> FixResult:
    """Accept all tracked changes in a DOCX (insertions kept, deletions removed)."""
    return await asyncio.to_thread(_accept_tracked_changes_sync, file_path, job_id)


def _accept_tracked_changes_sync(file_path: str, job_id: str) -> FixResult:

    doc = Document(file_path)
    body = doc.element.body
    accepted_ins = 0
    accepted_del = 0

    # Accept insertions: unwrap <w:ins> — keep child runs, remove wrapper
    for ins in body.findall(f".//{{{_WP_NS}}}ins"):
        parent = ins.getparent()
        if parent is None:
            continue
        idx = list(parent).index(ins)
        # Move children before removing wrapper
        for child in list(ins):
            ins.remove(child)
            parent.insert(idx, child)
            idx += 1
        parent.remove(ins)
        accepted_ins += 1

    # Accept deletions: remove <w:del> and all their contents entirely
    for deletion in body.findall(f".//{{{_WP_NS}}}del"):
        parent = deletion.getparent()
        if parent is not None:
            parent.remove(deletion)
            accepted_del += 1

    # Also remove revision properties (rPrChange, pPrChange, sectPrChange, tblPrChange)
    revision_tags = [
        "rPrChange",
        "pPrChange",
        "sectPrChange",
        "tblPrChange",
        "tblGridChange",
        "trPrChange",
        "tcPrChange",
    ]
    revision_removed = 0
    for tag in revision_tags:
        for elem in body.findall(f".//{{{_WP_NS}}}{tag}"):
            parent = elem.getparent()
            if parent is not None:
                parent.remove(elem)
                revision_removed += 1

    doc.save(file_path)

    total = accepted_ins + accepted_del + revision_removed
    return FixResult(
        tool_name="accept_tracked_changes",
        job_id=job_id,
        success=True,
        description=(
            f"Accepted tracked changes: {accepted_ins} insertion(s), "
            f"{accepted_del} deletion(s), {revision_removed} revision properties removed"
        ),
        after_value=f"{total} changes resolved",
    )


async def strip_hidden_text(
    file_path: str,
    job_id: str,
) -> FixResult:
    """Remove text runs marked as hidden (w:vanish) from a DOCX."""
    return await asyncio.to_thread(_strip_hidden_text_sync, file_path, job_id)


def _strip_hidden_text_sync(file_path: str, job_id: str) -> FixResult:

    doc = Document(file_path)
    removed = 0

    def _strip_from_paragraphs(paragraphs):
        nonlocal removed
        for para in paragraphs:
            for run in list(para.runs):
                rpr = run._element.find(f"{{{_WP_NS}}}rPr")
                if rpr is not None:
                    vanish = rpr.find(f"{{{_WP_NS}}}vanish")
                    if vanish is not None:
                        # Remove the entire run
                        run._element.getparent().remove(run._element)
                        removed += 1

    _strip_from_paragraphs(doc.paragraphs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _strip_from_paragraphs(cell.paragraphs)

    doc.save(file_path)

    return FixResult(
        tool_name="strip_hidden_text",
        job_id=job_id,
        success=True,
        description=f"Removed {removed} hidden text run(s)",
        after_value=f"{removed} removed",
    )


async def remove_empty_paragraphs(
    file_path: str,
    job_id: str,
    max_consecutive: int = 1,
) -> FixResult:
    """Collapse runs of empty paragraphs down to at most max_consecutive.

    An empty paragraph is one with no visible text and no embedded objects
    (images, page breaks, etc.). This is finer-grained than remove_blank_pages.
    """
    return await asyncio.to_thread(
        _remove_empty_paragraphs_sync,
        file_path,
        job_id,
        max_consecutive,
    )


def _remove_empty_paragraphs_sync(
    file_path: str,
    job_id: str,
    max_consecutive: int,
) -> FixResult:

    doc = Document(file_path)
    removed = 0
    consecutive_empty = 0

    for para in list(doc.paragraphs):
        if _is_empty_paragraph(para):
            consecutive_empty += 1
            if consecutive_empty > max_consecutive:
                parent = para._element.getparent()
                if parent is not None:
                    parent.remove(para._element)
                    removed += 1
        else:
            consecutive_empty = 0

    doc.save(file_path)

    return FixResult(
        tool_name="remove_empty_paragraphs",
        job_id=job_id,
        success=True,
        description=(f"Removed {removed} excess empty paragraph(s) (max consecutive: {max_consecutive})"),
        after_value=f"{removed} removed",
    )


def _is_empty_paragraph(para) -> bool:
    """Check if a paragraph is empty (no visible text, no page breaks, no images)."""
    if para.text.strip():
        return False

    # Check for embedded objects (images, breaks, etc.)
    elem = para._element
    # Page breaks
    if elem.findall(f".//{{{_WP_NS}}}br"):
        return False
    # Drawing elements (images)
    wp_ns = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    return not (elem.findall(f".//{{{wp_ns}}}inline") or elem.findall(f".//{{{wp_ns}}}anchor"))


# ── Style normalization ─────────────────────────────────────────────────

_HEADING_PREFIXES = ("heading", "title", "subtitle", "toc", "caption")


async def normalize_styles(
    file_path: str,
    job_id: str,
    target_body_font: str = "Calibri",
    target_body_size_pt: float = 11.0,
    normalize_line_spacing: bool = True,
) -> FixResult:
    """Normalize body text to a consistent style while preserving heading hierarchy.

    Forces body-text paragraphs to use a uniform font family, size, and
    line spacing. Headings (Heading 1-9, Title, Subtitle, TOC, Caption)
    are left unchanged to preserve document structure.
    """
    return await asyncio.to_thread(
        _normalize_styles_sync,
        file_path,
        job_id,
        target_body_font,
        target_body_size_pt,
        normalize_line_spacing,
    )


def _normalize_styles_sync(
    file_path: str,
    job_id: str,
    target_body_font: str,
    target_body_size_pt: float,
    normalize_line_spacing: bool,
) -> FixResult:

    doc = Document(file_path)
    changed_runs = 0
    changed_spacing = 0
    fonts_seen: set[str] = set()

    for para in doc.paragraphs:
        # Skip headings and structural styles
        style_name = (para.style.name or "").lower() if para.style else ""
        if any(style_name.startswith(prefix) for prefix in _HEADING_PREFIXES):
            continue

        # Normalize line spacing on body paragraphs
        if normalize_line_spacing:
            pf = para.paragraph_format
            if pf.line_spacing_rule is not None or pf.line_spacing is not None:
                # Only touch if explicitly set (don't force where default applies)
                from docx.enum.text import WD_LINE_SPACING  # noqa: PLC0415

                if pf.line_spacing_rule != WD_LINE_SPACING.MULTIPLE or pf.line_spacing != 1.15:
                    pf.line_spacing = 1.15
                    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                    changed_spacing += 1

        # Normalize font on each run
        for run in para.runs:
            did_change = False

            if run.font.name and run.font.name != target_body_font:
                fonts_seen.add(run.font.name)
                run.font.name = target_body_font
                did_change = True

            if run.font.size is not None:
                current_pt = run.font.size / 12700
                if abs(current_pt - target_body_size_pt) > 0.5:
                    run.font.size = Pt(target_body_size_pt)
                    did_change = True

            if did_change:
                changed_runs += 1

    # Also normalize table cell body text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        did_change = False
                        if run.font.name and run.font.name != target_body_font:
                            fonts_seen.add(run.font.name)
                            run.font.name = target_body_font
                            did_change = True
                        if run.font.size is not None:
                            current_pt = run.font.size / 12700
                            if abs(current_pt - target_body_size_pt) > 0.5:
                                run.font.size = Pt(target_body_size_pt)
                                did_change = True
                        if did_change:
                            changed_runs += 1

    doc.save(file_path)

    fonts_replaced = ", ".join(sorted(fonts_seen)[:5])
    if len(fonts_seen) > 5:
        fonts_replaced += f" (+{len(fonts_seen) - 5} more)"

    return FixResult(
        tool_name="normalize_styles",
        job_id=job_id,
        success=True,
        description=(
            f"Normalized {changed_runs} run(s) to {target_body_font} {target_body_size_pt}pt"
            + (f", {changed_spacing} paragraph(s) line spacing" if changed_spacing else "")
            + (f". Replaced fonts: {fonts_replaced}" if fonts_replaced else "")
        ),
        before_value=fonts_replaced or "mixed",
        after_value=f"{target_body_font} {target_body_size_pt}pt",
    )
