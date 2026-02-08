"""Page break fixes: remove consecutive breaks, strip manual breaks."""

from __future__ import annotations

import asyncio

from app.core.log import logger
from app.schema.fix import FixResult

__all__ = ("fix_page_breaks", "remove_manual_breaks")

_WP_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


async def fix_page_breaks(
    file_path: str,
    job_id: str,
    strategy: str = "remove_consecutive",
) -> FixResult:
    """Fix problematic page breaks. Strategies: remove_consecutive, remove_all."""
    if strategy == "remove_all":
        return await remove_manual_breaks(file_path, job_id)
    return await asyncio.to_thread(
        _remove_consecutive_breaks_sync, file_path, job_id,
    )


def _remove_consecutive_breaks_sync(file_path: str, job_id: str) -> FixResult:
    from docx import Document

    doc = Document(file_path)
    removed = 0
    prev_was_break = False

    for para in list(doc.paragraphs):
        is_break = _is_page_break_only(para)

        if is_break and prev_was_break:
            parent = para._element.getparent()
            if parent is not None:
                parent.remove(para._element)
                removed += 1
                continue

        prev_was_break = is_break

    doc.save(file_path)
    return FixResult(
        tool_name="fix_page_breaks",
        job_id=job_id,
        success=True,
        description=f"Removed {removed} consecutive page break(s)",
        after_value=f"{removed} removed",
    )


async def remove_manual_breaks(file_path: str, job_id: str) -> FixResult:
    """Remove all manual page breaks from a DOCX."""
    return await asyncio.to_thread(_remove_manual_breaks_sync, file_path, job_id)


def _remove_manual_breaks_sync(file_path: str, job_id: str) -> FixResult:
    from docx import Document

    doc = Document(file_path)
    removed = 0

    for para in doc.paragraphs:
        # Remove page-break-before from paragraph format
        if para.paragraph_format.page_break_before:
            para.paragraph_format.page_break_before = False
            removed += 1

        # Remove inline page break runs
        for run in para.runs:
            br_elements = run._element.findall(f"{{{_WP_NS}}}br")
            for br in br_elements:
                if br.get(f"{{{_WP_NS}}}type") == "page":
                    run._element.remove(br)
                    removed += 1

    doc.save(file_path)
    return FixResult(
        tool_name="remove_manual_breaks",
        job_id=job_id,
        success=True,
        description=f"Removed {removed} manual page break(s)",
        after_value=f"{removed} removed",
    )


def _is_page_break_only(para) -> bool:
    """Check if a paragraph is only a page break with no text content."""
    if para.text.strip():
        return False

    for run in para.runs:
        br_elements = run._element.findall(f"{{{_WP_NS}}}br")
        for br in br_elements:
            if br.get(f"{{{_WP_NS}}}type") == "page":
                return True

    if para.paragraph_format.page_break_before:
        return True

    return False
