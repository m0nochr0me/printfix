"""Structural analysis for DOCX files using python-docx."""

import asyncio
import re
import zipfile

from docx import Document
from docx.oxml.ns import qn

from app.core.config import settings
from app.core.log import logger
from app.schema.diagnosis import (
    DiagnosisIssue,
    IssueSeverity,
    IssueSource,
    IssueType,
)

__all__ = ("analyze_docx",)

_EMU_PER_INCH = 914400
_HALF_PT_PER_PT = 2


# Fonts commonly available on most print systems
_SAFE_FONTS = {
    "Arial",
    "Calibri",
    "Cambria",
    "Courier New",
    "Georgia",
    "Helvetica",
    "Tahoma",
    "Times New Roman",
    "Verdana",
    "Consolas",
    "Segoe UI",
    "Trebuchet MS",
    "Palatino Linotype",
    "Book Antiqua",
    "Garamond",
    "Century Gothic",
    "Lucida Sans",
    "Symbol",
    "Wingdings",
}


async def analyze_docx(file_path: str, job_id: str) -> list[DiagnosisIssue]:
    """Perform structural analysis on a DOCX file."""
    logger.info(f"Job {job_id}: running DOCX structural analysis on {file_path}")
    return await asyncio.to_thread(_analyze_docx_sync, file_path, job_id)


def _analyze_docx_sync(file_path: str, job_id: str) -> list[DiagnosisIssue]:

    issues: list[DiagnosisIssue] = []
    try:
        doc = Document(file_path)
        issues.extend(_check_margins(doc))
        issues.extend(_check_fonts(doc))
        issues.extend(_check_tables(doc))
        issues.extend(_check_images(doc))
        issues.extend(_check_paragraph_indents(doc))
        issues.extend(_check_page_breaks(doc))
        issues.extend(_check_hidden_content(doc))
        issues.extend(_check_columns(doc))
        issues.extend(_check_tracked_changes(file_path))
    except Exception:
        logger.exception(f"Job {job_id}: DOCX structural analysis failed")
    return issues


def _check_margins(doc) -> list[DiagnosisIssue]:
    """Check for inconsistent or too-small margins across sections."""
    issues: list[DiagnosisIssue] = []
    min_margin_emu = int(settings.DIAGNOSIS_MIN_MARGIN_INCHES * _EMU_PER_INCH)

    margin_sets: list[tuple[int, dict[str, int]]] = []
    for i, section in enumerate(doc.sections, 1):
        margins = {
            "top": section.top_margin or 0,
            "bottom": section.bottom_margin or 0,
            "left": section.left_margin or 0,
            "right": section.right_margin or 0,
        }
        margin_sets.append((i, margins))

        # Check for too-small margins
        for side, value in margins.items():
            if 0 < value < min_margin_emu:
                inches = value / _EMU_PER_INCH
                issues.append(
                    DiagnosisIssue(
                        type=IssueType.margin_violation,
                        severity=IssueSeverity.warning,
                        source=IssueSource.structural,
                        description=(
                            f'Section {i}: {side} margin is {inches:.2f}" '
                            f'(minimum {settings.DIAGNOSIS_MIN_MARGIN_INCHES}" recommended)'
                        ),
                        suggested_fix="set_margins",
                        confidence=0.9,
                    )
                )

    # Check for inconsistent margins across sections
    if len(margin_sets) > 1:
        first_margins = margin_sets[0][1]
        inconsistent = [sec_num for sec_num, margins in margin_sets[1:] if margins != first_margins]
        if inconsistent:
            issues.append(
                DiagnosisIssue(
                    type=IssueType.inconsistent_margins,
                    severity=IssueSeverity.warning,
                    source=IssueSource.structural,
                    description=(
                        f"Margins differ across sections: sections {inconsistent} have different margins than section 1"
                    ),
                    suggested_fix="set_margins",
                    confidence=0.85,
                )
            )

    return issues


def _check_fonts(doc) -> list[DiagnosisIssue]:
    """Check for uncommon fonts and small font sizes."""
    issues: list[DiagnosisIssue] = []
    min_pt = settings.DIAGNOSIS_MIN_FONT_PT

    fonts_used: set[str] = set()
    small_font_locations: list[str] = []

    # Check paragraphs
    for i, para in enumerate(doc.paragraphs, 1):
        for run in para.runs:
            if run.font.name:
                fonts_used.add(run.font.name)
            if run.font.size is not None:
                pt = run.font.size / 12700  # EMU to pt
                if pt < min_pt and run.text.strip():
                    small_font_locations.append(f"paragraph {i} ({pt:.1f}pt)")

    # Check table cells
    for t_idx, table in enumerate(doc.tables, 1):
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.font.name:
                            fonts_used.add(run.font.name)
                        if run.font.size is not None:
                            pt = run.font.size / 12700
                            if pt < min_pt and run.text.strip():
                                small_font_locations.append(f"table {t_idx} ({pt:.1f}pt)")

    # Flag uncommon fonts
    uncommon = fonts_used - _SAFE_FONTS
    for font in uncommon:
        issues = [
            *issues,
            DiagnosisIssue(
                type=IssueType.non_embedded_font,
                severity=IssueSeverity.warning,
                source=IssueSource.structural,
                description=f"Font '{font}' may not be available on print server",
                suggested_fix="replace_font",
                confidence=0.6,
            ),
        ]

    # Flag small fonts (aggregate)
    if small_font_locations:
        sample = small_font_locations[:5]
        issues.append(
            DiagnosisIssue(
                type=IssueType.small_font,
                severity=IssueSeverity.warning,
                source=IssueSource.structural,
                description=(
                    f"Small font sizes detected (< {min_pt}pt) in: "
                    + ", ".join(sample)
                    + (f" and {len(small_font_locations) - 5} more" if len(small_font_locations) > 5 else "")
                ),
                suggested_fix="adjust_font_size",
                confidence=0.85,
            )
        )

    return issues


def _check_tables(doc) -> list[DiagnosisIssue]:
    """Check if table widths exceed page width."""
    issues: list[DiagnosisIssue] = []

    # Get printable width from first section
    if not doc.sections:
        return issues
    section = doc.sections[0]
    page_width = section.page_width or 0
    left_margin = section.left_margin or 0
    right_margin = section.right_margin or 0
    printable_width = page_width - left_margin - right_margin

    if printable_width <= 0:
        return issues

    for t_idx, table in enumerate(doc.tables, 1):
        # Sum column widths from first row
        max_row_width = 0
        for row in table.rows:
            row_width = 0
            for cell in row.cells:
                cell_width = cell.width
                if cell_width:
                    row_width += cell_width
            max_row_width = max(max_row_width, row_width)

        if max_row_width > printable_width:
            overflow_pct = ((max_row_width - printable_width) / printable_width) * 100
            issues.append(
                DiagnosisIssue(
                    type=IssueType.table_overflow,
                    severity=IssueSeverity.critical,
                    source=IssueSource.structural,
                    description=(f"Table {t_idx} width exceeds printable area by {overflow_pct:.0f}%"),
                    suggested_fix="auto_fit_tables",
                    confidence=0.9,
                )
            )

    return issues


def _check_images(doc) -> list[DiagnosisIssue]:
    """Check if inline images exceed printable width."""
    issues: list[DiagnosisIssue] = []

    # Get printable width from first section
    if not doc.sections:
        return issues
    section = doc.sections[0]
    page_width = section.page_width or 0
    left_margin = section.left_margin or 0
    right_margin = section.right_margin or 0
    printable_width = page_width - left_margin - right_margin

    if printable_width <= 0:
        return issues

    printable_width_inches = printable_width / _EMU_PER_INCH
    overflow_images: list[tuple[int, float, float]] = []  # (para_num, img_width_in, overflow_pct)

    # Check inline images in paragraphs
    for para_num, para in enumerate(doc.paragraphs, 1):
        for run in para.runs:
            # Check for inline shapes (images)
            inline_shapes = run._element.findall(
                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
            )
            for drawing in inline_shapes:
                # Get extent (size) of the image
                extent = drawing.find(
                    ".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent"
                )
                if extent is not None:
                    cx = extent.get("cx")  # width in EMUs
                    if cx:
                        img_width = int(cx)
                        img_width_inches = img_width / _EMU_PER_INCH

                        # Calculate size relative to page
                        img_pct_of_page = (img_width / page_width) * 100 if page_width > 0 else 0

                        if img_width > printable_width:
                            overflow_pct = ((img_width - printable_width) / printable_width) * 100
                            overflow_images.append((para_num, img_width_inches, overflow_pct))

                            # Severity based on image size and overflow
                            if img_pct_of_page > 20 and overflow_pct > 10:
                                severity = IssueSeverity.critical
                                severity_note = "significant overflow"
                            elif img_pct_of_page > 5 and overflow_pct > 5:
                                severity = IssueSeverity.warning
                                severity_note = "moderate overflow"
                            else:
                                severity = IssueSeverity.info
                                severity_note = "minor overflow"

                            issues.append(
                                DiagnosisIssue(
                                    type=IssueType.image_overflow,
                                    severity=severity,
                                    source=IssueSource.structural,
                                    location=f"paragraph {para_num}",
                                    description=(
                                        f'Image width {img_width_inches:.2f}" exceeds printable area '
                                        f'{printable_width_inches:.2f}" by {overflow_pct:.0f}% '
                                        f"({severity_note})"
                                    ),
                                    suggested_fix="resize_image_to_fit",
                                    confidence=0.9,
                                )
                            )

    return issues


def _check_paragraph_indents(doc) -> list[DiagnosisIssue]:
    """Check for paragraph indents that consume excessive printable width."""
    issues: list[DiagnosisIssue] = []

    if not doc.sections:
        return issues
    section = doc.sections[0]
    page_width = section.page_width or 0
    left_margin = section.left_margin or 0
    right_margin = section.right_margin or 0
    printable_width = page_width - left_margin - right_margin

    if printable_width <= 0:
        return issues

    max_indent_emu = int(settings.DIAGNOSIS_MAX_INDENT_INCHES * _EMU_PER_INCH)
    excessive_locations: list[str] = []
    affected_paragraphs: list[dict] = []  # structured data for fix targeting
    max_left_seen = 0
    max_right_seen = 0
    has_critical = False

    def _scan_paragraph(para, label: str, index: int, context: str) -> None:
        nonlocal max_left_seen, max_right_seen, has_critical

        pf = para.paragraph_format
        left = pf.left_indent or 0
        right = pf.right_indent or 0
        first_line = pf.first_line_indent or 0

        max_left_seen = max(max_left_seen, left)
        max_right_seen = max(max_right_seen, right)

        loc_count_before = len(excessive_locations)
        already_flagged = False

        # Negative indents: text extends into/past margin toward page edge
        # left_margin + left_indent < 0 means text goes off the left edge
        if left < 0 or right < 0:
            left_in = left / _EMU_PER_INCH
            right_in = right / _EMU_PER_INCH
            # Critical if indent pushes text past the margin entirely
            if left < 0 and left_margin + left < 0:
                excessive_locations.append(f'{label} (L={left_in:+.2f}" overflows page edge)')
                has_critical = True
                already_flagged = True
            elif left < 0 and left_margin + left < min(left_margin // 2, max_indent_emu // 2):
                excessive_locations.append(f'{label} (L={left_in:+.2f}" into margin)')
                already_flagged = True
            if right < 0 and right_margin + right < 0:
                excessive_locations.append(f'{label} (R={right_in:+.2f}" overflows page edge)')
                has_critical = True
                already_flagged = True
            elif right < 0 and right_margin + right < min(right_margin // 2, max_indent_emu // 2):
                excessive_locations.append(f'{label} (R={right_in:+.2f}" into margin)')
                already_flagged = True

        # Also check negative first-line indent (hanging) that extends past margin
        if first_line < 0 and left + first_line < 0 and left_margin + left + first_line < 0:
            fl_in = first_line / _EMU_PER_INCH
            excessive_locations.append(f'{label} (hanging {fl_in:+.2f}" overflows page edge)')
            has_critical = True
            already_flagged = True

        # Positive excessive indents: wastes printable space
        if not already_flagged and (left > max_indent_emu or right > max_indent_emu):
            left_in = left / _EMU_PER_INCH
            right_in = right / _EMU_PER_INCH
            excessive_locations.append(f'{label} (L={left_in:.2f}" R={right_in:.2f}")')

        # Flag if combined positive indents consume >40% of printable width
        if not already_flagged:
            effective_left = left + max(first_line, 0)
            combined = effective_left + max(right, 0)
            if combined > printable_width * 0.4:
                if not any(label in loc for loc in excessive_locations):
                    combined_in = combined / _EMU_PER_INCH
                    excessive_locations.append(f'{label} (combined {combined_in:.2f}")')
                # Critical if <50% of page left for content
                if combined > printable_width * 0.5:
                    has_critical = True

        # Record structured data for any flagged paragraph
        if len(excessive_locations) > loc_count_before:
            affected_paragraphs.append(
                {
                    "index": index,
                    "context": context,
                    "left": round(left / _EMU_PER_INCH, 3),
                    "right": round(right / _EMU_PER_INCH, 3),
                    "first_line": round(first_line / _EMU_PER_INCH, 3),
                }
            )

    # Scan body paragraphs
    for i, para in enumerate(doc.paragraphs, 1):
        _scan_paragraph(para, f"paragraph {i}", index=i, context="body")

    # Scan table cells
    for t_idx, table in enumerate(doc.tables, 1):
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _scan_paragraph(para, f"table {t_idx}", index=t_idx, context=f"table_{t_idx}")

    if excessive_locations:
        sample = excessive_locations[:5]
        max_left_in = max_left_seen / _EMU_PER_INCH
        max_right_in = max_right_seen / _EMU_PER_INCH
        severity = IssueSeverity.critical if has_critical else IssueSeverity.warning

        issues.append(
            DiagnosisIssue(
                type=IssueType.inconsistent_indent,
                severity=severity,
                source=IssueSource.structural,
                description=(
                    f'Excessive paragraph indents detected (max L={max_left_in:.2f}"'
                    f' R={max_right_in:.2f}") in: '
                    + ", ".join(sample)
                    + (f" and {len(excessive_locations) - 5} more" if len(excessive_locations) > 5 else "")
                ),
                suggested_fix="adjust_paragraph_indents",
                confidence=0.85,
                metadata={"affected_paragraphs": affected_paragraphs},
            )
        )

    return issues


def _check_page_breaks(doc) -> list[DiagnosisIssue]:
    """Detect problematic page break patterns."""
    issues: list[DiagnosisIssue] = []
    consecutive_breaks = 0
    prev_was_break = False

    for i, para in enumerate(doc.paragraphs, 1):
        has_break = False

        # Check for explicit page break in runs
        for run in para.runs:
            if run._element.xml and "w:br" in run._element.xml:
                xml = run._element.xml
                if 'w:type="page"' in xml or "w:type='page'" in xml:
                    has_break = True
                    break

        # Check paragraph properties for page break before
        pf = para.paragraph_format
        if pf.page_break_before:
            has_break = True

        if has_break and prev_was_break:
            consecutive_breaks += 1
        elif has_break:
            consecutive_breaks = 1
        else:
            consecutive_breaks = 0

        prev_was_break = has_break

        if consecutive_breaks >= 2:
            issues.append(
                DiagnosisIssue(
                    type=IssueType.blank_page,
                    severity=IssueSeverity.warning,
                    source=IssueSource.structural,
                    location=f"near paragraph {i}",
                    description="Consecutive page breaks create blank page(s)",
                    suggested_fix="remove_blank_pages",
                    confidence=0.8,
                )
            )
            consecutive_breaks = 0  # don't double-report

    return issues


def _check_hidden_content(doc) -> list[DiagnosisIssue]:
    """Detect hidden text runs."""
    issues: list[DiagnosisIssue] = []
    hidden_count = 0

    for para in doc.paragraphs:
        for run in para.runs:
            rpr = run._element.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr")
            if rpr is not None:
                vanish = rpr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vanish")
                if vanish is not None and run.text.strip():
                    hidden_count += 1

    if hidden_count:
        issues.append(
            DiagnosisIssue(
                type=IssueType.hidden_content,
                severity=IssueSeverity.info,
                source=IssueSource.structural,
                description=f"{hidden_count} hidden text run(s) found in document",
                confidence=0.95,
            )
        )

    return issues


def _check_columns(doc) -> list[DiagnosisIssue]:
    """Check for multi-column layouts."""
    issues: list[DiagnosisIssue] = []

    for i, section in enumerate(doc.sections, 1):
        sectPr = section._sectPr
        cols = sectPr.xpath("./w:cols")
        if cols:
            num = cols[0].get(qn("w:num"))
            if num and int(num) > 1:
                issues.append(
                    DiagnosisIssue(
                        type=IssueType.multi_column_layout,
                        severity=IssueSeverity.info,
                        source=IssueSource.structural,
                        description=f"Section {i} uses a multi-column layout ({num} columns). This may cause layout issues during conversion.",
                        suggested_fix="set_columns",
                        metadata={"section_indices": [i], "current_columns": int(num)},
                    )
                )
    return issues


def _check_tracked_changes(file_path: str) -> list[DiagnosisIssue]:
    """Check for unresolved tracked changes via raw XML."""
    issues: list[DiagnosisIssue] = []

    try:
        with zipfile.ZipFile(file_path, "r") as z:
            if "word/document.xml" not in z.namelist():
                return issues
            xml_content = z.read("word/document.xml").decode("utf-8")

        # ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        insertions = len(re.findall("<w:ins\\b", xml_content))
        deletions = len(re.findall("<w:del\\b", xml_content))

        if insertions or deletions:
            issues.append(
                DiagnosisIssue(
                    type=IssueType.tracked_changes,
                    severity=IssueSeverity.warning,
                    source=IssueSource.structural,
                    description=(
                        f"Document contains unresolved tracked changes: "
                        f"{insertions} insertion(s), {deletions} deletion(s)"
                    ),
                    confidence=0.95,
                )
            )
    except Exception:
        logger.debug(f"Failed to check tracked changes in {file_path}")

    return issues
